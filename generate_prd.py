import os, sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_prd_documents():
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("Patrika Matrimony")
    r_title.bold = True
    r_title.font.size = Pt(26)
    r_title.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B) # Deep Red

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Product Requirements Document (Product / UX Specification)\nVersion M1.0 · Draft for Product Review · July 2026")
    r_sub.font.size = Pt(13)
    r_sub.italic = True
    r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    # Meta Table
    table_meta = doc.add_table(rows=5, cols=2)
    table_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Product Name", "Patrika Matrimony"),
        ("Brand Owner", "Rajasthan Patrika (Patrika Group)"),
        ("Target Platforms", "Mobile (Android Primary, iOS Compatible, Web Preview)"),
        ("Tagline / Positioning", "“Trusted matches from Rajasthan Patrika”"),
        ("Companion Specs", "PRD-Patrika-Matrimony.md / Expo React Native Source Repository")
    ]
    for idx, (k, v) in enumerate(meta_data):
        cell_k = table_meta.cell(idx, 0)
        cell_v = table_meta.cell(idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.8)
        
        p_k = cell_k.paragraphs[0]
        r_k = p_k.add_run(k)
        r_k.bold = True
        r_k.font.color.rgb = RGBColor(0x8E, 0x1B, 0x1B)
        
        p_v = cell_v.paragraphs[0]
        p_v.add_run(v)
        
        set_cell_background(cell_k, "FADBD8")
        set_cell_background(cell_v, "FDF2E9")
        set_cell_margins(cell_k, top=120, bottom=120, left=150, right=150)
        set_cell_margins(cell_v, top=120, bottom=120, left=150, right=150)

    doc.add_paragraph()

    # Helper for Headings
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(18)
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x96, 0x28, 0x1B)
        return h

    def add_h3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
        return h

    # Section 1
    add_h1("1. Executive Summary & Product Architecture")
    
    add_h2("1.1 Product Vision & Positioning")
    doc.add_paragraph(
        "Patrika Matrimony is a serious, trust-first mobile application designed for Rajasthan Patrika, "
        "one of India's leading newspaper groups renowned for its trusted matrimonial classifieds. The app bridges "
        "the physical print legacy of Rajasthan Patrika with modern digital matching, offering a respectful, family-inclusive, "
        "and culturally aligned matrimonial ecosystem."
    )
    doc.add_paragraph(
        "Unlike casual dating apps, Patrika Matrimony is built specifically for serious marriage seekers and parents/relatives. "
        "It emphasizes trust badges, verified community identities (Rajput, Agarwal, Brahmin, Marwari, Jain, Sikh, Sindhi, etc.), "
        "gotra/horoscope details, and seamless linking with Rajasthan Patrika newspaper classified ads."
    )

    add_h2("1.2 Target User Personas")
    personas = [
        ("The Self-Seeker (Individual)", "Young working professionals (aged 22–38) seeking a compatible partner while maintaining cultural and family alignment."),
        ("The Parent / Relative Initiator", "Parents, siblings, or elders creating and managing profiles for children or relatives. Requires clear typography, simple navigation, and multi-user privacy options."),
        ("The Print-to-Digital Reader", "Traditional readers transitioning from Rajasthan Patrika's Sunday matrimonial newspaper classifieds into the mobile app."),
        ("The Community & NRI Member", "Rajasthani diaspora across India (Mumbai, Bengaluru, Delhi) and abroad (USA, UK, UAE) seeking verified matches within specific ancestral origins.")
    ]
    for title, desc in personas:
        p = doc.add_paragraph()
        r = p.add_run(f"• {title}: ")
        r.bold = True
        p.add_run(desc)

    add_h2("1.3 Core Product Pillars")
    table_pillars = doc.add_table(rows=6, cols=2)
    table_pillars.alignment = WD_TABLE_ALIGNMENT.CENTER
    pillars = [
        ("Product Pillar", "What the User Experiences"),
        ("Print-to-Digital Ad Integration", "Allows users to search, locate, and link their Rajasthan Patrika print classified ad directly to their digital profile, unlocking the 'Patrika Verified' badge."),
        ("Trust & Multi-Level Verification", "7-step verification system including live selfie AI check, Govt ID (Aadhaar/PAN/Passport), Mobile OTP, Education, and Income badges."),
        ("Community & Cultural Precision", "Tailored matching on Gotra, Sub-Caste, Manglik status, Ancestral Native District, Horoscope/Nakshatra, and Diet preferences."),
        ("Privacy & Security First", "Photo privacy controls, phone number masking, anti-screenshot measures, blocking/reporting, and parent/relative profile management modes."),
        ("Freemium Monetization Model", "Free profile creation and browsing with contextual paywalls for Gold, Platinum, and Relationship Manager Assisted plans.")
    ]
    for idx, (col1, col2) in enumerate(pillars):
        c1 = table_pillars.cell(idx, 0)
        c2 = table_pillars.cell(idx, 1)
        c1.width = Inches(2.3)
        c2.width = Inches(4.7)
        
        p1 = c1.paragraphs[0]
        p2 = c2.paragraphs[0]
        
        r1 = p1.add_run(col1)
        r2 = p2.add_run(col2)
        
        if idx == 0:
            r1.bold = True
            r2.bold = True
            r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_background(c1, "C0392B")
            set_cell_background(c2, "C0392B")
        else:
            r1.bold = True
            set_cell_background(c1, "F9EBEA" if idx % 2 == 1 else "FFFFFF")
            set_cell_background(c2, "F9EBEA" if idx % 2 == 1 else "FFFFFF")
            
        set_cell_margins(c1, top=100, bottom=100, left=120, right=120)
        set_cell_margins(c2, top=100, bottom=100, left=120, right=120)

    doc.add_paragraph()

    # Section 2
    add_h1("2. Screen-by-Screen & Section-by-Section Product Specifications")

    screens = [
        {
            "num": "2.1",
            "title": "Splash & Boot Screen (`app/(auth)/splash.tsx`)",
            "purpose": "First brand impression, introducing Patrika Matrimony heritage and primary onboarding gateways.",
            "entry": "App launch / cold start.",
            "elements": [
                "Deep Red (#C0392B) to dark burgundy gradient canvas",
                "Animated Patrika Matrimony crest & bold brand typography",
                "Tagline: 'Trusted matches from Rajasthan Patrika'",
                "Primary CTAs: 'Create Profile' and 'Login'",
                "Community shortcut chips (Rajput, Marwari, Brahmin, Jain)",
                "Newspaper classified link banner at bottom"
            ],
            "rationale": "Establishes institutional trust instantly while giving fast access to both new registration and login flows."
        },
        {
            "num": "2.2",
            "title": "Login Screen (`app/(auth)/login.tsx`)",
            "purpose": "Dual-mode login supporting quick Mobile OTP and traditional Email/Password authentication.",
            "entry": "Tapping 'Login' on Splash screen.",
            "elements": [
                "Two main tab switchers: 'Mobile OTP' | 'Email & Password'",
                "Country code selector (+91 default) + 10-digit phone input",
                "'I am a Parent/Relative creating profile for family' toggle switch",
                "Primary red button: 'Get OTP'",
                "Forgot password link & help drawer trigger"
            ],
            "rationale": "Parents and older family members prefer phone OTP login without managing complex passwords."
        },
        {
            "num": "2.3",
            "title": "OTP Verification Screen (`app/(auth)/otp.tsx`)",
            "purpose": "Secure 6-digit OTP verification with resend timer.",
            "entry": "Submitting mobile number on Login screen.",
            "elements": [
                "6 individual digit input boxes with auto-focus advance",
                "30-second resend countdown timer",
                "Phone number edit link",
                "Primary CTA button: 'Verify & Proceed'",
                "Auto-dispatch demo user session (P001) for instant test entry"
            ],
            "rationale": "Eliminates friction during verification while enforcing mobile number ownership."
        },
        {
            "num": "2.4",
            "title": "13-Step Profile Creation Wizard (`app/(auth)/onboarding/step1.tsx` to `step13.tsx`)",
            "purpose": "Comprehensive step-by-step onboarding collecting essential matrimonial parameters.",
            "entry": "Tapping 'Create Profile' on Splash screen.",
            "elements": [
                "Step 1: Profile Creator (Self, Son, Daughter, Brother, Sister, Relative)",
                "Step 2: Mother Tongue Selection (Hindi, Marwari, Punjabi, Gujarati, etc.)",
                "Step 3: Basic Physical Details (Gender, DOB, Height, Physical Status, Marital Status)",
                "Step 4: Religion & Community (Religion, Caste, Sub-caste, Gotra, Manglik Status)",
                "Step 5: Location Details (Country, State, City with Rajasthan quick-chips)",
                "Step 6: Education & Career (Degree, Field, Employment Type, Occupation, Annual Income)",
                "Step 7: Family Background & Ancestral Origin (Family Status, Native District)",
                "Step 8: Lifestyle Preferences (Diet, Smoking, Drinking, Hobbies)",
                "Step 9: Horoscope Details Optional (Star/Nakshatra, Birth Time, Birth Place)",
                "Step 10: Institutional Details (College/University, Company Name, Current Role)",
                "Step 11: Account Setup (Full Name, Email, Password)",
                "Step 12: Photo Upload & Verification Prompt (Primary photo, gallery slots, selfie hint)",
                "Step 13: Partner Preferences (Age range, height range, caste, location, diet)"
            ],
            "rationale": "Chunking profile creation into 13 short steps prevents drop-off and yields complete, high-quality profile data."
        },
        {
            "num": "2.5",
            "title": "Home Screen (`app/(tabs)/home.tsx`)",
            "purpose": "Central match discovery hub featuring personalized recommendation rows and print ad integration.",
            "entry": "Default tab after login.",
            "elements": [
                "Top header: Patrika Matrimony red logo + Notification bell (badge count 3)",
                "Section 1: 'Recommended for You' — 10 horizontal scrolling profile cards matched to partner preferences",
                "Section 2: 'New Profiles' — profiles joined in last 30 days with 'New' badge",
                "Section 3: 'Verified Profiles' — profiles with verified badges",
                "Section 4: 'Nearby Profiles' — profiles in same city/state (e.g. Jaipur, Rajasthan)",
                "Section 5: 'From Rajasthan Patrika Ads' — print classified linked profiles with newspaper icon",
                "ProfileCard component with photo, name, age, city, caste, match %, and 3 action buttons (Interest, View, Shortlist)"
            ],
            "rationale": "Horizontal scrolling rows allow users to explore diverse match categories effortlessly."
        },
        {
            "num": "2.6",
            "title": "Search & Discovery Screen (`app/(tabs)/search.tsx`)",
            "purpose": "Advanced search engine with a 17-category filter drawer for pin-point match discovery.",
            "entry": "Search tab in bottom bar.",
            "elements": [
                "Top search input field + quick filter chips ('All', 'Nearby', 'New', 'Verified', 'Newspaper Ad')",
                "Filter Drawer with 17 collapsible filter sections: Religion, Caste, Gotra, Mother Tongue, City/State, Income, Education, Occupation, Diet, Manglik, Photo Only, Age Slider, Height Range",
                "Active filter counter badge on filter button",
                "Search results list: compact horizontal profile cards with instant filter matching"
            ],
            "rationale": "Matrimonial searches require exact parameter matching. The 17 filters empower families to find exact matches."
        },
        {
            "num": "2.7",
            "title": "Profile Details Screen (`app/profile/[id].tsx`)",
            "purpose": "Comprehensive view of a candidate's full profile, photos, horoscope, and contact actions.",
            "entry": "Tapping any profile card across the app.",
            "elements": [
                "Hero photo section with full-width image, gradient overlay, gallery dots, and fullscreen viewer",
                "Name, Age, Match % badge, City/State, Marital status, Caste, and Verification badges",
                "Action bar: Shortlist (heart), Send Interest (handshake), Message (chat), Call (phone)",
                "Contextual paywall trigger: Free members tapping Message/Call see 'Upgrade to Gold' bottom sheet",
                "Detailed sections: About Me, Basic Details, Religious & Social Background (Gotra, Manglik), Professional Details, Family Background, Horoscope Accordion, Hobbies Chips, Photos Gallery, and Similar Profiles horizontal carousel",
                "Sticky bottom bar: Send Interest & Message buttons"
            ],
            "rationale": "Presents all sensitive and cultural information in a clean, dignified, expandable layout."
        },
        {
            "num": "2.8",
            "title": "Chats & Calls Screen (`app/(tabs)/chats.tsx`)",
            "purpose": "Communication hub for active conversations, interest responses, and voice call history.",
            "entry": "Chats tab in bottom bar.",
            "elements": [
                "Top tab switcher: 'Chats' | 'Calls'",
                "Sub-filters: 'All', 'Accepted', 'New Interests'",
                "Chat list items: Profile photo, Name, Last message text, Timestamp, Unread count badge",
                "Calls tab: Incoming, Outgoing, and Missed call history with quick callback button",
                "Free Plan banner: 'Upgrade to Gold to chat with matches →'"
            ],
            "rationale": "Separating chats and call history keeps communications organized while serving as a primary upgrade trigger."
        },
        {
            "num": "2.9",
            "title": "Chat Conversation Screen (`app/chat/[id].tsx`)",
            "purpose": "Real-time 1-on-1 messaging interface with candidate or family contact.",
            "entry": "Tapping any chat item in Chats list.",
            "elements": [
                "Header: Back arrow, candidate photo, name, online status, voice call & video call buttons",
                "Message list: Sent messages in red (#C0392B) right bubbles, received messages in white left bubbles",
                "Read status tick marks (double tick)",
                "Input area with text input, attachment icon, and send button",
                "Paywall lock for Free plan users with upgrade bottom sheet"
            ],
            "rationale": "Clean, secure messaging environment with read receipts and instant upgrade options."
        },
        {
            "num": "2.10",
            "title": "Shortlist & Interests Screen (`app/(tabs)/shortlist.tsx`)",
            "purpose": "Manage saved profiles and track sent/received interest invitations.",
            "entry": "Shortlist tab in bottom bar.",
            "elements": [
                "2-column grid of shortlisted profile cards with quick remove option",
                "Interests Received section: Profile cards with 'Accept' (green) and 'Decline' (gray) buttons",
                "Interests Sent section: Status chips ('Interest Sent', 'Accepted', 'Declined')",
                "Empty state illustrations for shortlist and interests"
            ],
            "rationale": "Allows users and parents to track shortlists and manage pending marriage proposals in one clean workspace."
        },
        {
            "num": "2.11",
            "title": "Rajasthan Patrika Newspaper Ads Screen (`app/newspaper-ads/index.tsx`)",
            "purpose": "Bridge traditional print classified ads with digital app profiles.",
            "entry": "Navigating from Home banner, Profile menu, or Newspaper Ad filter.",
            "elements": [
                "Header: '🗞️ Rajasthan Patrika Matrimonial Ads'",
                "Banner: 'Have an ad in Rajasthan Patrika? Link your print ad to get Patrika Verified badge'",
                "Filter chips: All Ads, Linked, Linkable + Community category chips (Hindu-Rajput, Jain, etc.)",
                "Print-style ad cards displaying edition date, city, ad text body, masked phone number, and linked profile card",
                "'Link My Ad' Modal: Edition date picker, city picker, Ad Reference ID input, search button, and confirmation link CTA"
            ],
            "rationale": "Creates a unique competitive advantage by digitizing Rajasthan Patrika's print advertiser base."
        },
        {
            "num": "2.12",
            "title": "Trust & Verification Centre Screen (`app/verification/index.tsx`)",
            "purpose": "Increase platform safety and trust by driving users to verify their identities.",
            "entry": "Profile tab ➔ Verification Centre.",
            "elements": [
                "Trust Score progress bar (e.g., 75% Verified)",
                "7 Verification Cards: Mobile OTP (Verified), Live Selfie AI Check (Pending/Verify), Govt ID Aadhaar/PAN (Upload), Education Certificate (Upload), Salary Slip/ITR (Upload), Location (NRI), LinkedIn/Social Link",
                "Interactive selfie camera capture simulation and document file picker",
                "Verification badge descriptions and status chips"
            ],
            "rationale": "Verified profiles receive 5x more responses. Gamifying verification builds platform credibility."
        },
        {
            "num": "2.13",
            "title": "Subscription & Monetization Screen (`app/subscription/index.tsx`)",
            "purpose": "Present premium membership plans and execute seamless upgrade checkout.",
            "entry": "Upgrade buttons across app, Profile tab, or paywall bottom sheets.",
            "elements": [
                "4 Plan Cards: Free (Gray), Gold ₹1,999/3mo (Gold), Platinum ₹3,999/6mo (Purple - Most Popular), Assisted ₹9,999/yr (Red)",
                "Feature comparison matrix table",
                "Payment Checkout Modal: Order summary, tax breakdown, payment mode selector (UPI: GPay/PhonePe/Paytm/BHIM, Credit/Debit Card, Net Banking)",
                "Payment Success Screen with green checkmark animation, transaction ref ID, and instant plan activation"
            ],
            "rationale": "Contextual, transparent pricing with Indian payment methods maximizes subscription conversion."
        },
        {
            "num": "2.14",
            "title": "My Profile & Account Screen (`app/(tabs)/profile.tsx`)",
            "purpose": "Manage personal profile, account settings, privacy, and membership status.",
            "entry": "Profile tab in bottom bar.",
            "elements": [
                "Profile Header: Photo, Name, Age, Membership Badge ('Free Member', 'Gold', 'Platinum'), Upgrade CTA",
                "Profile completion progress bar (75% complete) with fill prompt",
                "Quick stats row: Interests Received (12), Accepted (5), Profile Views (234)",
                "Menu sections: My Account (Edit Profile modal, Partner Preferences, Subscription, Verification), Newspaper Integration, Privacy & Safety, Support, Logout"
            ],
            "rationale": "Central control panel for managing visibility, editing details, and monitoring response metrics."
        }
    ]

    for sc in screens:
        add_h2(f"{sc['num']} {sc['title']}")
        p_purp = doc.add_paragraph()
        r_p = p_purp.add_run("Purpose: ")
        r_p.bold = True
        p_purp.add_run(sc["purpose"])
        
        p_ent = doc.add_paragraph()
        r_e = p_ent.add_run("Entry Point / Trigger: ")
        r_e.bold = True
        p_ent.add_run(sc["entry"])
        
        doc.add_paragraph("Key Screen Elements & Functional Specifications:").runs[0].bold = True
        for elem in sc["elements"]:
            doc.add_paragraph(f"  • {elem}")
            
        p_rat = doc.add_paragraph()
        r_r = p_rat.add_run("Product Rationale: ")
        r_r.bold = True
        p_rat.add_run(sc["rationale"])

    # Section 3
    add_h1("3. Monetization & Pricing Tiers Specification")
    doc.add_paragraph(
        "Patrika Matrimony employs a freemium model. Free users can register, build complete profiles, "
        "browse matches, and send limited interests. Access to contact numbers, direct chat messaging, "
        "horoscope matching details, and dedicated relationship management requires an upgraded plan."
    )
    
    table_plans = doc.add_table(rows=5, cols=5)
    table_plans.alignment = WD_TABLE_ALIGNMENT.CENTER
    plan_headers = ["Plan Name", "Price & Duration", "Target User", "Key Unlocked Features", "Limitations"]
    plan_rows = [
        ("Free", "₹0 / Lifetime", "New registrants, browsers", "Create profile, 20 profile views/day, 5 interests/day, basic search filters", "No phone numbers, no chat, no horoscope match"),
        ("Gold", "₹1,999 / 3 Months", "Active marriage seekers", "Unlimited profile views, unlimited interests, view phone/email, direct in-app messaging, priority search", "No voice calls, no relationship manager"),
        ("Platinum (Popular)", "₹3,999 / 6 Months", "High-intent families", "Everything in Gold + in-app voice calls, advanced horoscope filter, private photo access, 10 profile boosts/mo", "No dedicated relationship manager"),
        ("Assisted", "₹9,999 / 1 Year", "Hectic professionals, parents", "Everything in Platinum + Dedicated Relationship Manager (RM), RM shortlists matches, offline meeting coordination, 24/7 priority support", "None")
    ]
    for c_idx, h_text in enumerate(plan_headers):
        cell = table_plans.cell(0, c_idx)
        p = cell.paragraphs[0]
        r = p.add_run(h_text)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "C0392B")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    for r_idx, row_data in enumerate(plan_rows, start=1):
        for c_idx, val in enumerate(row_data):
            cell = table_plans.cell(r_idx, c_idx)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if c_idx == 0:
                r.bold = True
            set_cell_background(cell, "F9EBEA" if r_idx % 2 == 1 else "FFFFFF")
            set_cell_margins(cell, top=100, bottom=100, left=100, right=100)

    doc.add_paragraph()

    # Section 4
    add_h1("4. Technical Build & System Architecture Summary")
    arch_points = [
        ("Frontend Stack", "React Native 0.86, Expo SDK 57, Expo Router v57 (file-based navigation), TypeScript ~6.0."),
        ("UI & Styling", "Custom StyleSheet tokens, React Native Paper UI elements, @expo/vector-icons, expo-linear-gradient."),
        ("State Management", "React Context API (`AppContext`) + `useReducer` with `@react-native-async-storage/async-storage` state persistence."),
        ("Data Layer", "100% local mock JSON datasets containing 200 profiles (`profiles.json`), 500 interests (`interests.json`), 100 chat threads (`chats.json`), 60 subscriptions (`subscriptions.json`), and 50 print ads (`newspaper-ads.json`)."),
        ("Build & Deployment", "Configured for local Expo Go testing, React Native Web bundling (`react-native-web`), and EAS Cloud APK generation (`eas.json`).")
    ]
    for title, desc in arch_points:
        p = doc.add_paragraph()
        r = p.add_run(f"• {title}: ")
        r.bold = True
        p.add_run(desc)

    doc.save(r'C:\Users\Anubhav.Shubham\Documents\Codex\Patrika Matrimony App\PatrikaMatrimony\PRD-Patrika-Matrimony.docx')
    print("DOCX generated successfully.")

create_prd_documents()
